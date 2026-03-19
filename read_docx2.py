import docx
import sys

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            full_text.append(' | '.join(row_data))
            
    return '\n'.join(full_text)

if __name__ == '__main__':
    print(read_docx(sys.argv[1]))
